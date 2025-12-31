# app/ingest_nextcloud.py
import os
import time
import json
import requests
import xml.etree.ElementTree as ET
import logging
import sys
import io
import tempfile
import shutil
from urllib.parse import urljoin, urlparse
from typing import Optional, List, Tuple, Dict, Set, Any
import hashlib
import urllib3
from concurrent.futures import ThreadPoolExecutor, as_completed

# --- Text Extraction Libraries ---
try:
    import fitz  # PyMuPDF (Handles PDF & EPUB fast)
    from docx import Document as DocxDocument
    import mobi
    import html2text
    import openpyxl
    import pandas as pd
    import mutagen
    from mutagen.id3 import ID3
    from mutagen.mp3 import MP3
    from mutagen.flac import FLAC
    from mutagen.mp4 import MP4
    from mutagen.oggvorbis import OggVorbis
except ImportError as e:
    print(f"ERROR: Missing required ingestion library: {e}")
    sys.exit(1)

# Optional: Check for cryptography
try:
    import cryptography
except ImportError:
    pass 

# Suppress SSL Warnings
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# LangChain and Chroma Imports
try:
    from langchain_core.documents import Document
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_chroma import Chroma
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError as e:
    print(f"CRITICAL: Missing AI dependencies: {e}")
    sys.exit(1)

# --- Configuration ---
if os.getenv("DOCKER_ENV") != "1" and os.path.exists(".env"):
    from dotenv import load_dotenv
    load_dotenv(".env")

NEXTCLOUD_URL = os.getenv("NEXTCLOUD_URL")
NEXTCLOUD_USER = os.getenv("NEXTCLOUD_USER")
NEXTCLOUD_PASS = os.getenv("NEXTCLOUD_PASS")
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
DEFAULT_MODEL = os.getenv("DEFAULT_MODEL", "qwen3:latest")

CHROMA_DIR = os.getenv("CHROMA_PERSIST_DIR", "/data/chroma_db")
INGESTION_TEMP_DIR = os.getenv("INGESTION_TEMP_DIR", "/data/nextcloud_temp")
LOG_FILE_PATH = "/data/nextcloud_ingest.log"
TOC_CACHE_FILE = "/data/nextcloud_toc_cache.json"

# Logic Configuration
PERSIST_FREQUENCY = 50
MAX_WORKERS = 8
TOC_CACHE_TTL = 86400 
MIN_TEXT_LENGTH = 10 

# File Categorization
SPREADSHEET_EXTS = [".csv", ".xlsx", ".json", ".xls"]
BOOK_EXTS = [".epub", ".mobi", ".pdf"] 
TEXT_EXTS = [".txt", ".md", ".docx"]
MEDIA_EXTS = [".mp3", ".m4b", ".mp4", ".mkv", ".avi", ".flac", ".wav", ".mov", ".webm", ".ogg"]

BLACKLIST_DIRS = [
    "Music", "Videos", "Photos", "Audio", "Thumbnails", "Preview", "Cache", "AppData", "Templates"
]
BLACKLIST_FULL_PATHS = ["Books/Audio"]

# --- Logging Setup ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] [%(name)s] %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(LOG_FILE_PATH)
    ],
    force=True
)
logger = logging.getLogger("NextcloudSync")

if not NEXTCLOUD_URL or not NEXTCLOUD_USER or not NEXTCLOUD_PASS:
    logger.error("Nextcloud settings missing (NEXTCLOUD_URL, USER, or PASS)")
    sys.exit(1)

NAMESPACES = {"d": "DAV:"}

# ----------------------
# AI Helper Functions
# ----------------------

def generate_summary(text: str) -> str:
    """Generates a high-level summary using the local LLM."""
    if not OLLAMA_URL: return ""
    trunc_text = text[:10000] # Safe limit
    prompt = f"Summarize the following document in 3-4 sentences. Capture the main topic, key themes, and purpose.\n\nDocument:\n{trunc_text}\n\nSummary:"
    
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/generate",
            json={"model": DEFAULT_MODEL, "prompt": prompt, "stream": False, "options": {"temperature": 0.0}},
            timeout=60
        )
        if resp.status_code == 200:
            return resp.json().get("response", "").strip()
    except Exception as e:
        logger.warning(f"Summary generation failed: {e}")
    return ""

def extract_spreadsheet_schema(file_path: str, filename: str) -> str:
    """Reads a spreadsheet and creates a rich text description."""
    try:
        df = None
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".csv":
            df = pd.read_csv(file_path, nrows=5)
            total_rows = sum(1 for _ in open(file_path, 'r', encoding='utf-8', errors='ignore')) - 1
        elif ext in [".xlsx", ".xls"]:
            df = pd.read_excel(file_path, nrows=5)
            total_rows = "Unknown (Excel)" 
        elif ext == ".json":
            df = pd.read_json(file_path)
            total_rows = len(df)
            df = df.head(5)

        if df is None: return ""

        columns = ", ".join(df.columns.tolist())
        sample = df.to_string(index=False, max_rows=3)
        
        description = (
            f"DATA FILE: {filename}\n"
            f"TYPE: Spreadsheet/Structured Data\n"
            f"COLUMNS: {columns}\n"
            f"TOTAL ROWS: ~{total_rows}\n"
            f"SAMPLE DATA:\n{sample}\n"
            f"USAGE: Use this file for budget, finance, or data analysis queries. "
            f"File available in Nextcloud at original path."
        )
        return description

    except Exception as e:
        logger.warning(f"Failed to extract schema for {filename}: {e}")
        return f"DATA FILE: {filename}. Could not parse structure automatically."

# ----------------------
# Multimedia Metadata Extraction (NEW)
# ----------------------

def extract_media_metadata(file_path: str) -> Dict[str, Any]:
    """Extracts technical and tag metadata from audio/video files using mutagen."""
    meta = {
        "title": None, "artist": None, "album": None, "year": None, "genre": None,
        "duration": 0, "bitrate": 0, "sample_rate": 0, "encoding": "unknown"
    }
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        audio = mutagen.File(file_path)
        if audio is None: return meta
        
        # Technical Stats
        if hasattr(audio, 'info'):
            meta["duration"] = getattr(audio.info, 'length', 0)
            meta["bitrate"] = getattr(audio.info, 'bitrate', 0)
            meta["sample_rate"] = getattr(audio.info, 'sample_rate', 0)
            
            # Simple encoding detection
            if ext == ".mp3": meta["encoding"] = "mp3"
            elif ext == ".flac": meta["encoding"] = "flac"
            elif ext == ".ogg": meta["encoding"] = "ogg"
            elif ext in [".m4a", ".mp4"]: meta["encoding"] = "aac/mp4"
            else: meta["encoding"] = type(audio.info).__name__

        # ID3 / Tags
        # Mutagen abstracts most tags into a dictionary-like interface
        tags = audio.tags
        if tags:
            if ext == ".mp3":
                try:
                    from mutagen.easyid3 import EasyID3
                    easy_tags = EasyID3(file_path)
                    meta["title"] = easy_tags.get('title', [None])[0]
                    meta["artist"] = easy_tags.get('artist', [None])[0]
                    meta["album"] = easy_tags.get('album', [None])[0]
                    meta["year"] = easy_tags.get('date', [None])[0]
                    meta["genre"] = easy_tags.get('genre', [None])[0]
                except Exception as e:
                    logger.debug(f"EasyID3 failed, falling back to raw ID3: {e}")
                    # Raw ID3 fallback
                    meta["title"] = str(tags.get('TIT2')) if 'TIT2' in tags else None
                    meta["artist"] = str(tags.get('TPE1')) if 'TPE1' in tags else None
                    meta["album"] = str(tags.get('TALB')) if 'TALB' in tags else None
                    meta["year"] = str(tags.get('TDRC', tags.get('TYER'))) if ('TDRC' in tags or 'TYER' in tags) else None
                    meta["genre"] = str(tags.get('TCON')) if 'TCON' in tags else None
            else:
                # Vorbis/FLAC/MP4 use names
                meta["title"] = tags.get('title', [None])[0] if tags.get('title') else None
                meta["artist"] = tags.get('artist', [None])[0] if tags.get('artist') else None
                meta["album"] = tags.get('album', [None])[0] if tags.get('album') else None
                meta["year"] = tags.get('date', [None])[0] if tags.get('date') else None
                meta["genre"] = tags.get('genre', [None])[0] if tags.get('genre') else None

    except Exception as e:
        logger.warning(f"Metadata extraction failed for {file_path}: {e}")
        
    return {k: v for k, v in meta.items() if v is not None} # Remove empty values

# ----------------------
# Text & Book Extraction (UPDATED)
# ----------------------

def format_toc(toc_list: list) -> str:
    """Converts PyMuPDF TOC list [[lvl, title, page], ...] into a readable string."""
    if not toc_list: return ""
    output = ["--- TABLE OF CONTENTS ---"]
    for item in toc_list:
        # Handle 3 or 4 item tuples (lvl, title, page, dest_dict)
        lvl, title, page = item[0], item[1], item[2]
        indent = "  " * (lvl - 1)
        output.append(f"{indent}- {title} (Pg {page})")
    return "\n".join(output)

def extract_text_from_mobi(file_path: str) -> Optional[str]:
    """Legacy helper for MOBI files."""
    tempdir = None
    try:
        tempdir, extracted_path = mobi.extract(file_path)
        if extracted_path:
            with open(extracted_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return html2text.html2text(content)
        return None
    except Exception as e:
        logger.error(f"Failed to extract text from MOBI at {file_path}: {e}")
        return None
    finally:
        if tempdir and os.path.exists(tempdir):
            shutil.rmtree(tempdir)

def extract_text_from_xlsx(file_path: str) -> Optional[str]:
    # Fallback for text extraction if schema fails or if simpler text needed
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        all_text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            all_text.append(f"--- Sheet: {sheet_name} ---")
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join([str(cell) for cell in row if cell is not None])
                if row_text.strip():
                    all_text.append(row_text)
        return "\n".join(all_text)
    except Exception as e:
        logger.error(f"Failed to extract text from XLSX at {file_path}: {e}")
        return None

def extract_content_and_meta(file_path: str) -> Tuple[Optional[str], Dict[str, Any]]:
    """
    Universal extractor. Returns (full_text, metadata_dict).
    Metadata dict contains 'toc' (string) and 'book_meta' (dict) for books.
    """
    ext = os.path.splitext(file_path)[1].lower().strip()
    if not ext: return None, {}
    
    extra_data = {"toc": None, "book_meta": {}}

    try:
        # --- FAST PATH: PyMuPDF (PDF & EPUB) ---
        if ext in [".pdf", ".epub"]:
            text = ""
            with fitz.open(file_path) as doc:
                # 1. Extract TOC
                try:
                    toc_list = doc.get_toc()
                    if toc_list:
                        extra_data["toc"] = format_toc(toc_list)
                except Exception as e:
                    logger.warning(f"TOC extraction failed for {file_path}: {e}")

                # 2. Extract Metadata (Title/Author)
                try:
                    extra_data["book_meta"] = doc.metadata
                except: pass

                # 3. Extract Text
                for page in doc:
                    text += page.get_text()
            return text, extra_data

        # --- STANDARD PATH: Other Formats ---
        elif ext in [".txt", ".md", ".csv", ".json"]:
            try:
                with open(file_path, "r", encoding="utf-8") as f: return f.read(), extra_data
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1", errors="ignore") as f: return f.read(), extra_data
        
        # --- DOCX ---
        elif ext == ".docx":
            document = DocxDocument(file_path)
            content_pieces = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip(): content_pieces.append(paragraph.text)
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text: content_pieces.append(row_text)
            return "\n".join(content_pieces), extra_data
            
        elif ext == ".xlsx":
            return extract_text_from_xlsx(file_path), extra_data
            
        elif ext == ".mobi":
            return extract_text_from_mobi(file_path), extra_data
            
    except Exception as e:
        err_str = str(e).lower()
        if "cryptography" in err_str:
             logger.warning(f"Missing 'cryptography' library for PDF: {file_path}")
        else:
             logger.error(f"Extraction error on {file_path}: {e}")
    
    return None, {}

# ----------------------
# Parallel Scanning Logic
# ----------------------

def get_file_category(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext in SPREADSHEET_EXTS: return "spreadsheet"
    if ext in BOOK_EXTS: return "book"
    if ext in TEXT_EXTS: return "text"
    if ext in MEDIA_EXTS: return "media"
    return "unknown"

def scan_single_folder(current_url: str) -> Tuple[List[str], Dict[str, Dict]]:
    """
    Worker function to scan a single folder. 
    Returns: (list_of_subdirs, dict_of_files_found)
    """
    subdirs = []
    files_found = {}
    
    propfind_data = """<?xml version="1.0" encoding="utf-8" ?>
    <d:propfind xmlns:d="DAV:">
      <d:prop>
        <d:getetag/>
        <d:resourcetype/>
        <d:getcontentlength/>
      </d:prop>
    </d:propfind>"""

    for attempt in range(3):
        try:
            resp = requests.request(
                "PROPFIND", current_url, 
                auth=(NEXTCLOUD_USER, NEXTCLOUD_PASS), 
                headers={"Depth": "1", "Content-Type": "application/xml"}, 
                data=propfind_data,
                verify=False, timeout=900
            )
            
            if resp.status_code == 404: return [], {}
            
            if resp.status_code == 207:
                root = ET.fromstring(resp.content)
                for r in root.findall("d:response", NAMESPACES):
                    href = r.find("d:href", NAMESPACES).text
                    prop = r.find("d:propstat/d:prop", NAMESPACES)
                    if not href or not prop: continue
                    
                    etag = prop.find("d:getetag", NAMESPACES).text.strip('"') if prop.find("d:getetag", NAMESPACES) is not None else "unknown"
                    
                    size = "0"
                    size_node = prop.find("d:getcontentlength", NAMESPACES)
                    if size_node is not None:
                        size = size_node.text

                    href_decoded = requests.utils.unquote(href)
                    prefix = f"/remote.php/dav/files/{NEXTCLOUD_USER}/"
                    
                    if prefix not in href_decoded: continue
                    rel_path = href_decoded.split(prefix, 1)[1]
                    item_url = urljoin(NEXTCLOUD_URL, href)
                    
                    if item_url.rstrip('/') == current_url.rstrip('/'): continue

                    is_collection = prop.find("d:resourcetype/d:collection", NAMESPACES) is not None
                    
                    if is_collection:
                        dir_name = rel_path.strip("/").split("/")[-1]
                        if rel_path.strip("/") not in BLACKLIST_FULL_PATHS and dir_name not in BLACKLIST_DIRS:
                            if not item_url.endswith("/"): item_url += "/"
                            subdirs.append(item_url)
                    else:
                        cat = get_file_category(rel_path)
                        if cat != "unknown":
                            files_found[rel_path] = {
                                "etag": etag, 
                                "category": cat,
                                "size": size
                            }
                
                return subdirs, files_found
            
            time.sleep(5)
        except Exception as e:
            logger.warning(f"Scan error {current_url}: {e}")
            time.sleep(5)
            
    return [], {}

def list_files_parallel(start_url: str, nc_state: Dict[str, Dict]):
    logger.info(f"Starting Parallel WebDAV scan ({MAX_WORKERS} threads)...")
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_url = {executor.submit(scan_single_folder, start_url): start_url}
        while future_to_url:
            done = [f for f in future_to_url if f.done()]
            for f in done:
                url = future_to_url.pop(f)
                try:
                    subdirs, files = f.result()
                    nc_state.update(files)
                    for sub in subdirs:
                        future_to_url[executor.submit(scan_single_folder, sub)] = sub
                except: pass
            if not done: time.sleep(0.1)

# ----------------------
# Sync Logic (UPDATED)
# ----------------------
def load_toc_cache() -> Optional[Dict[str, Dict]]:
    if not os.path.exists(TOC_CACHE_FILE): return None
    try:
        with open(TOC_CACHE_FILE, 'r') as f:
            data = json.load(f)
        if time.time() - data.get('timestamp', 0) > TOC_CACHE_TTL: return None
        return data.get('files', {})
    except: return None

def save_toc_cache(files: Dict[str, Dict]):
    try:
        data = {"timestamp": time.time(), "files": files}
        with open(TOC_CACHE_FILE, 'w') as f: json.dump(data, f)
    except: pass

def get_db_state(vectordb, target_source: Optional[str] = None) -> Dict[str, Dict]:
    logger.info("Fetching database state...")
    try:
        # Optimization: If targeted, only fetch that specific document
        if target_source:
             results = vectordb.get(where={"source": target_source}, include=["metadatas"])
        else:
             # Warning: This is heavy for large DBs. Consider pagination in future.
             results = vectordb.get(include=["metadatas"])
             
        state = {}
        if results and results.get("metadatas"):
            for meta in results["metadatas"]:
                if meta and "source" in meta:
                    state[meta["source"]] = {
                        "etag": meta.get("etag", "unknown"),
                        "type": meta.get("type", "unknown")
                    }
        return state
    except Exception as e:
        logger.warning(f"Failed to fetch DB state: {e}")
        return {}

def sync_nextcloud_files(target_rel_path: Optional[str] = None):
    if not os.path.exists(INGESTION_TEMP_DIR): os.makedirs(INGESTION_TEMP_DIR)
    
    logger.info(f"--- Starting Nextcloud Ingestion (Target: {target_rel_path or 'Entire Library'}) ---")
    
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectordb = Chroma(collection_name="nextcloud_docs", embedding_function=embeddings, persist_directory=CHROMA_DIR)

    # Optimization: Pass target to get_db_state to allow filtered fetch
    db_state = get_db_state(vectordb, target_source=target_rel_path)
    
    # If a specific path is provided, we skip the full scan and just probe that item
    if target_rel_path:
        logger.info(f"Targeted ingestion requested for: {target_rel_path}")
        root_url = urljoin(NEXTCLOUD_URL, f"/remote.php/dav/files/{NEXTCLOUD_USER}/")
        target_url = urljoin(root_url, requests.utils.quote(target_rel_path))
        
        nc_state = {}
        
        # 1. Try treating it as a SINGLE FILE first (Naive Probe)
        logger.info(f"Probing {target_url} as single file (Depth 0)...")
        try:
             resp = requests.request("PROPFIND", target_url, auth=(NEXTCLOUD_USER, NEXTCLOUD_PASS), headers={"Depth": "0"}, timeout=30, verify=False)
             logger.info(f"Probe Status: {resp.status_code}")
             
             if resp.status_code == 207:
                  root = ET.fromstring(resp.content)
                  res = root.find("d:response", NAMESPACES)
                  prop = res.find("d:propstat/d:prop", NAMESPACES) if res is not None else None
                  
                  is_collection = False
                  if prop is not None:
                       if prop.find("d:resourcetype/d:collection", NAMESPACES) is not None:
                            is_collection = True
                  
                  if not is_collection:
                       cat = get_file_category(target_rel_path)
                       if cat != "unknown" and prop:
                           etag = prop.find("d:getetag", NAMESPACES).text.strip('"') if prop.find("d:getetag", NAMESPACES) is not None else "unknown"
                           size = prop.find("d:getcontentlength", NAMESPACES).text if prop.find("d:getcontentlength", NAMESPACES) is not None else "0"
                           nc_state[target_rel_path] = {"etag": etag, "category": cat, "size": size}
                           logger.info(f"Single file identified. Etag: {etag}, Size: {size}")
                       else:
                           logger.warning(f"File found but category '{cat}' unknown or props missing.")
                  else:
                       logger.info("Target is a directory (Collection). Switching to recursive scan.")
                       if not target_url.endswith("/"): target_url += "/"
                       list_files_parallel(target_url, nc_state)
             else:
                  logger.warning(f"Probe failed {resp.status_code}. Trying recursive scan as fallback...")
                  if not target_url.endswith("/"): target_url += "/"
                  list_files_parallel(target_url, nc_state)
        except Exception as e:
             logger.error(f"Targeted probe error: {e}")
    else:
        nc_state = load_toc_cache()
        if not nc_state:
            nc_state = {}
            root_url = urljoin(NEXTCLOUD_URL, f"/remote.php/dav/files/{NEXTCLOUD_USER}/")
            list_files_parallel(root_url, nc_state)
            save_toc_cache(nc_state)
    
    to_ingest = []
    # ONLY remove files if we did a full scan (nc_state is the whole library)
    if not target_rel_path:
        files_to_remove = [path for path in db_state if path not in nc_state]
        if files_to_remove:
            for path in files_to_remove:
                try: vectordb._collection.delete(where={"source": path})
                except: pass
    else:
        # If targeted, we don't delete anything, we just update/add the target.
        pass

    # --- RETROFIT LOGIC ---
    for path, info in nc_state.items():
        db_record = db_state.get(path)
        should_process = False
        
        if target_rel_path:
             should_process = True # Force update if targeted
        elif not db_record:
            should_process = True
        elif db_record["etag"] != info["etag"]:
            should_process = True
        else:
            current_type = db_record["type"]
            expected_category = info["category"]
            
            # Retrofit: Spreadsheets
            if expected_category == "spreadsheet" and current_type != "spreadsheet_metadata":
                should_process = True
            
            # Retrofit: Books
            elif expected_category == "book" and (current_type != "book_card" and current_type != "book_summary"):
                 should_process = True

        if should_process:
            to_ingest.append((path, info["category"], info["etag"], info.get("size", "0")))

    logger.info(f"Sync Plan: {len(to_ingest)} files to ingest/update.")

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=200, length_function=len)
    processed_count = 0
    
    for i, (rel_path, category, etag, size) in enumerate(to_ingest, start=1):
        try:
            vectordb._collection.delete(where={"source": rel_path})
        except: pass

        doc_id_base = hashlib.sha256(rel_path.encode()).hexdigest()
        fname = os.path.basename(rel_path)
        common_metadata = {
            "source": rel_path, "filename": fname, "etag": etag, "size": int(size),
            "source_server": NEXTCLOUD_URL, "source_user": NEXTCLOUD_USER
        }

        # --- ROUTE 1: SPREADSHEETS ---
        if category == "spreadsheet":
            logger.info(f"Processing SPREADSHEET Schema: {fname}")
            temp_path = os.path.join(INGESTION_TEMP_DIR, f"{doc_id_base}_{fname}")
            encoded_path = requests.utils.quote(rel_path)
            file_url = urljoin(NEXTCLOUD_URL, f"/remote.php/dav/files/{NEXTCLOUD_USER}/{encoded_path}")
            
            try:
                with requests.get(file_url, auth=(NEXTCLOUD_USER, NEXTCLOUD_PASS), stream=True, verify=False, timeout=600) as r:
                    r.raise_for_status()
                    with open(temp_path, 'wb') as f:
                        for chunk in r.iter_content(chunk_size=8192): f.write(chunk)
                
                schema = extract_spreadsheet_schema(temp_path, fname)
                vectordb.add_documents([Document(
                    page_content=schema, 
                    metadata={**common_metadata, "type": "spreadsheet_metadata", "is_remote": True}
                )])
                processed_count += 1
            except Exception as e:
                logger.error(f"Error processing spreadsheet {fname}: {e}")
            finally:
                if os.path.exists(temp_path): os.remove(temp_path)
            continue

        # --- ROUTE 2: MEDIA ---
        if category == "media":
            logger.info(f"Processing MEDIA Metadata: {fname}")
            temp_path = os.path.join(INGESTION_TEMP_DIR, f"{doc_id_base}_{fname}")
            encoded_path = requests.utils.quote(rel_path)
            file_url = urljoin(NEXTCLOUD_URL, f"/remote.php/dav/files/{NEXTCLOUD_USER}/{encoded_path}")
            
            try:
                # Download small chunk for metadata if possible, but mutagen often needs the whole file for ID3/info
                # We'll download the file like books/spreadsheets.
                with requests.get(file_url, auth=(NEXTCLOUD_USER, NEXTCLOUD_PASS), stream=True, verify=False, timeout=600) as r:
                    r.raise_for_status()
                    with open(temp_path, 'wb') as f:
                        for chunk in r.iter_content(chunk_size=8192): f.write(chunk)
                
                media_meta = extract_media_metadata(temp_path)
                
                # Build a rich descriptive string for the page_content
                # EMPHASIS: Repeat key terms to boost semantic ranking against large texts (Aristotle)
                desc_parts = [f"Media File: {fname}", f"Path: {rel_path}"]
                
                title = media_meta.get("title")
                artist = media_meta.get("artist")
                album = media_meta.get("album")
                
                if title:
                    desc_parts.append(f"Title: {title}")
                    desc_parts.append(f"Song Name: {title}") # Synonym for retrieval
                
                if artist:
                    desc_parts.append(f"Artist: {artist}")
                    desc_parts.append(f"Performer: {artist}") # Synonym
                    
                if album:
                     desc_parts.append(f"Album: {album}")
                
                # Semantic boosters
                if title and artist:
                     desc_parts.append(f"Music Track: {title} by {artist}")
                     desc_parts.append(f"Audio Track: {title} by {artist}")

                
                tech_details = []
                if media_meta.get("encoding"): tech_details.append(f"Encoding: {media_meta['encoding']}")
                if media_meta.get("bitrate"): tech_details.append(f"Bitrate: {int(media_meta['bitrate']/1000)}kbps")
                if media_meta.get("duration"): 
                    m, s = divmod(int(media_meta['duration']), 60)
                    tech_details.append(f"Duration: {m}:{s:02d}")
                
                if tech_details:
                    desc_parts.append(f"Technical: {', '.join(tech_details)}")
                
                desc_parts.append(f"Size: {size} bytes")
                
                meta_doc = Document(
                    page_content="\n".join(desc_parts),
                    metadata={**common_metadata, **media_meta, "type": "media", "category": "audio_video", "chunk": 0}
                )
                vectordb.add_documents([meta_doc])
                processed_count += 1
            except Exception as e:
                logger.error(f"Error processing media {fname}: {e}")
            finally:
                if os.path.exists(temp_path): os.remove(temp_path)
            continue

        # --- ROUTE 3: TEXT & BOOKS (UPDATED) ---
        if category == "text" or category == "book":
            logger.info(f"[{i}/{len(to_ingest)}] Processing: {rel_path} ({size} bytes)")
            encoded_path = requests.utils.quote(rel_path)
            file_url = urljoin(NEXTCLOUD_URL, f"/remote.php/dav/files/{NEXTCLOUD_USER}/{encoded_path}")
            temp_path = os.path.join(INGESTION_TEMP_DIR, f"{doc_id_base}_{fname}")
            
            try:
                with requests.get(file_url, auth=(NEXTCLOUD_USER, NEXTCLOUD_PASS), stream=True, verify=False, timeout=900) as r:
                    r.raise_for_status()
                    with open(temp_path, 'wb') as f:
                        for chunk in r.iter_content(chunk_size=8192): f.write(chunk)
                
                # CALL UPDATED EXTRACTOR
                content, extra_data = extract_content_and_meta(temp_path)
                
                if content and len(content.strip()) > MIN_TEXT_LENGTH:
                    docs_to_add = []
                    
                    # --- A. BOOK FEATURES ---
                    if category == "book":
                        # 1. LIBRARY CARD
                        book_meta = extra_data.get("book_meta", {})
                        summary = generate_summary(content)
                        
                        card_content = (
                            f"LIBRARY CARD:\n"
                            f"Title: {book_meta.get('title', fname)}\n"
                            f"Author: {book_meta.get('author', 'Unknown')}\n"
                            f"Format: {book_meta.get('format', 'PDF/EPUB')}\n"
                            f"Filename: {fname}\n"
                            f"Synopsis: {summary}"
                        )
                        docs_to_add.append(Document(
                            page_content=card_content,
                            metadata={**common_metadata, "type": "book_card"} # New specialized type
                        ))

                        # 2. TABLE OF CONTENTS
                        if extra_data.get("toc"):
                            docs_to_add.append(Document(
                                page_content=extra_data["toc"],
                                metadata={**common_metadata, "type": "book_toc"}
                            ))

                    # --- B. STANDARD CHUNKS ---
                    chunks = text_splitter.split_text(content)
                    for ix, c in enumerate(chunks):
                        docs_to_add.append(Document(
                            page_content=c, 
                            metadata={**common_metadata, "type": category, "chunk_index": ix}
                        ))
                    
                    vectordb.add_documents(docs_to_add)
                    processed_count += 1
                    logger.info(f"-> Ingested {len(chunks)} chunks + Metadata.")
            
            except Exception as e:
                logger.error(f"Error processing {rel_path}: {e}")
            finally:
                if os.path.exists(temp_path): os.remove(temp_path)
        
        # Periodic Persist
        if processed_count % PERSIST_FREQUENCY == 0:
             if hasattr(vectordb, 'persist'): vectordb.persist()

    if hasattr(vectordb, 'persist'): vectordb.persist()
    logger.info(f"Sync finished.")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="NextCloud Ingestion Service")
    parser.add_argument("--path", type=str, help="Specific file or folder path to ingest (relative to user root)")
    args = parser.parse_args()
    
    sync_nextcloud_files(target_rel_path=args.path)
